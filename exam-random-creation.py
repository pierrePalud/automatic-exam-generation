import pandas as pd
import os

from docx import Document
from docx.shared import Inches
from docx.shared import Pt

from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows


def generate_one_test(df, num_questions_per_level):
    list_df_questions = []
        
    for diff_level in df['DifficultyFromQuestioner'].unique():    
        if num_questions_per_level[diff_level] > 0:
            df_quest_diff = df[df['DifficultyFromQuestioner'] == diff_level].sample(n=num_questions_per_level[diff_level])        
            list_df_questions.append(df_quest_diff)

    df_test = pd.concat(list_df_questions)
    return df_test


def create_word_doc(df_test, id_test, exam_title, exam_version=True):
    document = Document()
    document.add_heading(f'{exam_title} - Student Version {id_test}', 0)

    p = document.add_paragraph('This test is just a ')
    p.add_run('Proof of Concept').bold = True
    p.add_run(' of automaticly and randomly generated exams.')

    document.add_heading('Instructions', level=1)
    document.add_paragraph('Answer each question with a short sentence. Exam duration : 15 minutes')

    document.add_heading('Questions', level=1)
    
    df_test = df_test.reset_index(drop=True)
    for index, row in df_test.iterrows():
        p = document.add_paragraph()
        if exam_version:
            font = p.add_run(
                f"{index + 1} - {row['Question']}\n"
            ).font
        else:
            font = p.add_run(
                f"{index + 1} - {row['Question']}\n\t{row['Answer']}"
            ).font
        
        font.size = Pt(10)

    if exam_version:
        if not(os.path.isdir("./students_exam")):
            os.mkdir("./students_exam")
        document.save(f'./students_exam/student_test_{id_test}.docx')
    else:
        if not(os.path.isdir("./corrections")):
            os.mkdir("./corrections")
        document.save(f'./corrections/correction_test_{id_test}.docx')
    


def save_pool_with_exams(df):
    workbook = Workbook()
    sheet = workbook.active

    for row in dataframe_to_rows(df, index=False, header=True):
        sheet.append(row)

    sheet.auto_filter.ref = sheet.dimensions
    
    if not(os.path.isdir("./questions")):
        os.mkdir("./questions")
    workbook.save("./questions/whole_exam_correction.xlsx")
    
    
def read_question_pool():
    workbook = load_workbook(filename="./questions_pool.xlsx")
    sheet = workbook.active

    data = sheet.values

    # Set the first row as the columns for the DataFrame
    cols = next(data)
    data = list(data)

    df = pd.DataFrame(data, columns=cols)
    return df

    
    
if __name__ == '__main__':
	try:
	    df = read_question_pool()
	    valid = True
	except:
		k = input("The questions_pool.xlsx has not been found. Put it in the same folder as this exe file. Press any key to exit.")
		valid = False

	if valid:
		exam_title = input('What title do you want for the exam ? (ex: Marketing Quiz 1) ')
		while len(exam_title) == 0:
			exam_title = input('Please enter a title (ex: Marketing Quiz 1) ')
	    
		num_distinct_exams = '1'
		while not(isinstance(num_distinct_exams, int)) or num_distinct_exams <= 0:
			num_distinct_exams = input('How many distinct exams do you want ? ')
			try:
				num_distinct_exams = int(num_distinct_exams)
				if num_distinct_exams <= 0:
					a = 1/0
			except:
				print(f"{num_distinct_exams} is not a valid number. Please enter a strictly positive integer (1, 2, 3, etc.).")

		num_questions_per_level = {}
		for diff_level in df['DifficultyFromQuestioner'].unique():
			num_questions_test = '1'
			num_pool = len(df.loc[df['DifficultyFromQuestioner'] == diff_level])

			while not(isinstance(num_questions_test, int)) or num_questions_test < 0 or num_questions_test > num_pool:
				num_questions_test = input(f'How many {diff_level.upper()} questions do you want per exam ? (max possible : {num_pool}) ')
				try:
					num_questions_test = int(num_questions_test)
					if num_questions_test < 0 or num_questions_test > num_pool:
						a = 1/0
				except:
					if num_questions_test > num_pool:
						print(f"You asked for {num_questions_test} per exam, but there are only {num_pool} {diff_level.upper()} questions in your question pool. Please enter a smaller number.")
					else:
						print(f"{num_questions_test} is not a valid number. Please enter a positive integer (0, 1, 10, 30, etc.).")

			num_questions_per_level[diff_level] = num_questions_test


		for id_test in range(1, num_distinct_exams+1):
			df_test = generate_one_test(df, num_questions_per_level)
			create_word_doc(df_test, id_test, exam_title, exam_version=True)
			create_word_doc(df_test, id_test, exam_title, exam_version=False)

			for idx_test, idx_df in enumerate(df_test.index): 
				df.at[idx_df, f'test_{id_test}'] = idx_test+1

		save_pool_with_exams(df)
